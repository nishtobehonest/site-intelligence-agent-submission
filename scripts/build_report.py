from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x16, 0x48, 0x7B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Site Intelligence Agent")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
title.paragraph_format.space_before = Pt(24)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Phase 1 Case Analysis & Progress Report")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x16, 0x48, 0x7B)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Site Intelligence Agent  |  April 2026").font.size = Pt(10)

add_hr(doc)

# ── Section 1 ──
h1(doc, "1. What Is This Project?")
body(doc, (
    "Imagine you are an HVAC technician standing in front of a broken rooftop air conditioning unit. "
    "You need to know how to safely shut off the power before touching any wires, but you do not have the manual with you. "
    "Normally you would call the office and wait."
))
body(doc, (
    "This system changes that. You type your question on your phone or tablet, and within seconds you get a cited, "
    "reliable answer pulled directly from official safety documents and equipment manuals."
))
body(doc, (
    "The core innovation is what happens when the system is not sure. Most AI tools will confidently give you a wrong answer. "
    "This system is designed to say 'I am not sure' or 'call your supervisor' when it does not have enough information. "
    "In a safety-critical environment, knowing when NOT to answer is more important than always sounding confident."
))

add_hr(doc)

# ── Section 2 ──
h1(doc, "2. How It Works: The Simple Version")
body(doc, "Think of the system as a very smart librarian with three bookshelves:")
bullet(doc, "OSHA safety regulations (government rules about workplace safety)", "Bookshelf 1 - Safety Rules: ")
bullet(doc, "Equipment manuals from manufacturers like Carrier, Trane, and Lennox", "Bookshelf 2 - Equipment Manuals: ")
bullet(doc, "50 job history records covering past repair and maintenance visits", "Bookshelf 3 - Past Job Records: ")
doc.add_paragraph()
body(doc, "When you ask a question, the system does four things in sequence:")
bullet(doc, "Searches all three bookshelves for the most relevant pages.", "Step 1 - Search: ")
bullet(doc, "Scores how confident it is in what it found: HIGH, PARTIAL, or LOW.", "Step 2 - Grade the match: ")
bullet(doc, "If confident enough, hands those pages to an AI (Gemini) and says: answer this using only what is on these pages.", "Step 3 - Generate answer: ")
bullet(doc, "Returns the answer with a trust label so the technician knows how much to rely on it.", "Step 4 - Label and return: ")

add_hr(doc)

# ── Section 3 ──
h1(doc, "3. What Was Phase 1?")
body(doc, (
    "Before Phase 1, the system was like a library with empty shelves. All the software was written and working, "
    "but there were no documents in it. Every question returned LOW confidence because there was nothing to search."
))
body(doc, "Phase 1 had one goal: stock the shelves and prove the three confidence levels work end-to-end with real data.")

h2(doc, "3.1  Starting State")
add_table(doc,
    ["Component", "State Before Phase 1"],
    [
        ["All pipeline code (retrieval, scoring, routing, LLM)", "Complete, written but untested with real data"],
        ["OSHA documents", "Empty folder (README placeholder only)"],
        ["Equipment manuals", "Empty folder (README placeholder only)"],
        ["Job history records", "Empty folder (README placeholder only)"],
        ["Chroma vector database", "Did not exist"],
        ["Evaluation test sets", "Stub files with 2 examples each"],
    ],
    col_widths=[3.0, 3.3]
)

h2(doc, "3.2  What I Did")
bullet(doc, "Downloaded 2 real OSHA PDFs: 29 CFR 1910.147 (Lockout/Tagout procedures) and 29 CFR 1910.303 (Electrical Safety requirements).", "OSHA documents: ")
bullet(doc, (
    "Downloaded 4 real equipment manuals: two versions of the Carrier 48LC rooftop unit (2017 and 2023 editions), "
    "Lennox SL280 gas furnace, and Trane XR15 heat pump. "
    "Having two Carrier versions is deliberate -- they list different refrigerant pressures, enabling contradiction testing in Phase 2."
), "Equipment manuals: ")
bullet(doc, (
    "Manually created 50 synthetic job history records covering 5 equipment types "
    "(Carrier, Trane, Lennox, York, Daikin) and 5 job types (preventive maintenance, emergency repair, "
    "inspection, installation, compliance check)."
), "Job history: ")
bullet(doc, (
    "Ran the ingestion pipeline -- all documents were split into 500-character chunks, "
    "converted into embedding vectors using the sentence-transformers/all-MiniLM-L6-v2 model, "
    "and stored in three separate Chroma collections."
), "Ingestion: ")

h2(doc, "3.3  Issues Found and Fixed")
add_table(doc,
    ["Issue", "Root Cause", "Fix Applied"],
    [
        [
            "Broken imports on startup",
            "LangChain reorganized its module structure in newer versions. "
            "langchain.text_splitter and langchain.schema.Document both moved to new packages.",
            "Updated imports in src/ingest.py to langchain_text_splitters and langchain_core.documents"
        ],
        [
            "All queries routing LOW (score 0.36 for a strong match)",
            "Wrong similarity formula. The code used 1 - L2_distance. "
            "For unit-normalized sentence-transformer vectors, the correct formula is 1 - (L2 squared / 2). "
            "This caused a genuine 0.80 match to appear as 0.36.",
            "Fixed formula in src/retriever.py. Lockout/tagout query now correctly scores 0.80 and routes HIGH."
        ],
        [
            "LLM calls failing with 404 errors",
            "Gemini 1.5 Pro and Gemini 2.0 Flash are deprecated and not available to new API keys.",
            "Updated default model to gemini-2.5-flash in src/llm.py"
        ],
        [
            "1-page placeholder PDFs in data folders",
            "Original files were AI-generated summaries, not real documents. Low page count means very few chunks and poor retrieval coverage.",
            "Replaced with real multi-page PDFs from official OSHA and manufacturer sources"
        ],
    ],
    col_widths=[1.6, 2.4, 2.3]
)

add_hr(doc)

# ── Section 4 ──
h1(doc, "4. What Is in the Database Now")
body(doc, "After ingestion, the Chroma vector database contains three collections with a total of 1,836 searchable chunks:")
add_table(doc,
    ["Collection", "Source Files", "Chunks", "Contents"],
    [
        ["osha", "2 PDFs", "283", "1910.147 Lockout/Tagout + 1910.303 Electrical Safety"],
        ["manuals", "4 PDFs", "1,372", "Carrier 48LC (2017 + 2023), Lennox SL280, Trane XR15"],
        ["job_history", "1 JSON file", "181", "50 synthetic HVAC job history records"],
        ["TOTAL", "", "1,836", ""],
    ],
    col_widths=[1.3, 1.2, 0.9, 2.9]
)
body(doc, (
    "A chunk is a small piece of a document, roughly 500 characters long. "
    "Each chunk is stored as a vector of 384 numbers that capture the meaning of the text, not just the keywords. "
    "This is why the system can match 'how do I turn off power safely' to a chunk that says "
    "'lockout/tagout energy control procedure' -- the meanings are close even though the words differ."
))

add_hr(doc)

# ── Section 5 ──
h1(doc, "5. Smoke Test Results")
body(doc, (
    "The smoke test runs three preset queries, one designed to trigger each confidence level. "
    "This verifies the full pipeline works end-to-end before moving to formal evaluation."
))

add_table(doc,
    ["Demo", "Query (abbreviated)", "Expected Route", "Actual Route", "Score", "Result"],
    [
        ["1 - Safety procedure", "Lockout/tagout for electrical HVAC maintenance", "HIGH", "HIGH", "0.80", "PASS"],
        ["2 - Spec lookup", "Refrigerant charge pressure for Carrier rooftop unit", "PARTIAL", "PARTIAL", "0.52", "PASS"],
        ["3 - Missing record", "Last maintenance on equipment ID HVAC-SITE07-UNIT-99", "LOW", "PARTIAL", "0.62", "PARTIAL PASS"],
    ],
    col_widths=[1.3, 2.1, 1.0, 1.0, 0.6, 0.9]
)

h2(doc, "Demo 1 - HIGH Confidence (Full Pass)")
body(doc, (
    "The lockout/tagout query found the OSHA 1910.147 document with a similarity score of 0.80. "
    "The system passed the retrieved context to Gemini and returned a cited answer. "
    "This is the ideal path: the technician gets a reliable, sourced answer immediately with no warnings."
))

h2(doc, "Demo 2 - PARTIAL Confidence (Full Pass)")
body(doc, (
    "The refrigerant pressure query scored 0.52, placing it in the middle zone. "
    "The system returned an answer but flagged it with a yellow warning: "
    "This answer is based on partial context. Verify before acting. "
    "The technician is informed rather than misled. "
    "True conflict detection between the 2017 and 2023 Carrier manuals is a Phase 2 enhancement."
))

h2(doc, "Demo 3 - LOW Confidence (Partial Pass)")
body(doc, (
    "The fabricated equipment ID query (HVAC-SITE07-UNIT-99) routed PARTIAL instead of LOW. "
    "The system still set Escalate=True and warned the technician to verify before acting, "
    "so the safety outcome is correct. However, the ideal behavior is a hard LOW with no LLM call."
))
body(doc, (
    "Root cause: the system uses semantic similarity. The query mentions 'HVAC maintenance,' "
    "which is semantically close to real job records -- even though the specific equipment ID does not exist. "
    "The system cannot distinguish 'this topic exists in the corpus' from 'this specific ID exists in the corpus.' "
    "This requires a targeted fix in Phase 2."
))

add_hr(doc)

# ── Section 6 ──
h1(doc, "6. What Works and What Does Not")

h2(doc, "Working Correctly")
bullet(doc, "End-to-end pipeline: question in, cited answer out.")
bullet(doc, "Correct cosine similarity scoring after formula fix.")
bullet(doc, "HIGH confidence path: retrieves, generates, cites source.")
bullet(doc, "PARTIAL confidence path: generates with explicit uncertainty warning.")
bullet(doc, "Gemini 2.5 Flash LLM integration.")
bullet(doc, "Three separate Chroma collections correctly isolated.")
bullet(doc, "Escalate flag fires correctly on both PARTIAL and LOW routes.")

h2(doc, "Not Yet Working")
bullet(doc, (
    "Demo 3 routes PARTIAL instead of LOW for fabricated equipment IDs. "
    "Semantic similarity cannot detect that a specific ID does not exist."
), "Equipment-ID-specific escalation: ")
bullet(doc, (
    "Two Carrier manual versions are in the same collection. "
    "The conflict detection function only triggers when results come from different collection names, "
    "not different files within the same collection."
), "Within-collection conflict detection: ")
bullet(doc, (
    "Eval files are still 2-entry stubs. "
    "Hallucination rate, coverage rate, and conflict detection rate cannot be measured yet."
), "Evaluation metrics: ")

add_hr(doc)

# ── Section 7 ──
h1(doc, "7. What Comes Next: Phase 2")
add_table(doc,
    ["Task", "Owner", "Why It Matters"],
    [
        [
            "Fix detect_conflicts to catch within-collection version conflicts",
            "Nishchay",
            "Enables true Demo 2 conflict surfacing between 2017 vs 2023 Carrier manuals"
        ],
        [
            "Fix Demo 3 LOW routing for fabricated equipment IDs",
            "Nishchay",
            "Ensures hard escalation when a job record truly does not exist"
        ],
        [
            "Expand ground_truth.json to 50 Q&A pairs",
            "Nishchay",
            "Required to measure coverage rate and hallucination rate"
        ],
        [
            "Expand adversarial.json to 20 out-of-corpus queries",
            "Nishchay",
            "Required to measure escalation rate"
        ],
        [
            "Expand contradictions.json to 15 conflict scenarios",
            "Nishchay",
            "Required to measure conflict detection rate"
        ],
        [
            "Run eval/run_eval.py and record baseline metrics",
            "Nishchay",
            "Establishes whether targets are met before Phase 3 threshold tuning"
        ],
    ],
    col_widths=[2.5, 1.0, 2.8]
)

add_hr(doc)

# ── Section 8 ──
h1(doc, "8. Metric Targets")
body(doc, "These targets were set at project kick-off and remain the Phase 5 success criteria:")
add_table(doc,
    ["Metric", "Target", "Current Status"],
    [
        ["Hallucination rate", "< 2%", "Not yet measured -- eval set needed"],
        ["Coverage rate", "> 80%", "Not yet measured -- eval set needed"],
        ["Escalation rate", "10-25%", "Not yet measured -- eval set needed"],
        ["Conflict detection rate", "Majority of contradiction scenarios", "Not working correctly yet"],
        ["Time to answer", "< 8 seconds", "Observed approx. 3-5 seconds in smoke test"],
    ],
    col_widths=[2.2, 1.2, 2.9]
)

add_hr(doc)

# ── Footer ──
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run("Site Intelligence Agent  |  April 2026")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save("Docs/Phase1_Case_Analysis.docx")
print("Saved: Docs/Phase1_Case_Analysis.docx")
