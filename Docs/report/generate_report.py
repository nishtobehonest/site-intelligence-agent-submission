"""
generate_report.py
------------------
Generates site_intelligence_report.docx from report_content.md using python-docx.

Run:
    pip install python-docx
    python Docs/report/generate_report.py

Output: Docs/report/site_intelligence_report.docx
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).parent
SOURCE_MD  = SCRIPT_DIR / "report_content.md"
OUTPUT_DOC = SCRIPT_DIR / "site_intelligence_report.docx"

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

COLOR_H1        = RGBColor(0x1F, 0x49, 0x7D)   # dark navy
COLOR_H2        = RGBColor(0x2E, 0x74, 0xB5)   # mid blue
COLOR_H3        = RGBColor(0x26, 0x4F, 0x78)   # steel blue
COLOR_CODE_BG   = RGBColor(0xF2, 0xF2, 0xF2)   # light grey for code cells
COLOR_SCSHOT_BG = RGBColor(0xE8, 0xF0, 0xFE)   # light blue placeholder
COLOR_CALLOUT   = RGBColor(0xFF, 0xF3, 0xCD)   # amber for blockquotes
COLOR_TABLE_HDR = RGBColor(0x1F, 0x49, 0x7D)   # navy header row

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def setup_document() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name  = "Calibri"
    h1.font.size  = Pt(18)
    h1.font.bold  = True
    h1.font.color.rgb = COLOR_H1
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name  = "Calibri"
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = COLOR_H2
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name  = "Calibri"
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = COLOR_H3
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)

    return doc


# ---------------------------------------------------------------------------
# XML helpers for shading and borders
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour (e.g. '1F497D')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_para_bg(para, hex_color: str):
    """Set a paragraph background/shading (for code blocks)."""
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _add_para_border(para, hex_color: str, border_size: int = 4):
    """Add a left border to a paragraph (blockquote / callout style)."""
    pPr     = para._p.get_or_add_pPr()
    pBdr    = OxmlElement("w:pBdr")
    left    = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(border_size * 4))  # eighths of a point
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_full_border(para, hex_color: str = "AAAAAA"):
    """Add a box border around a paragraph (placeholder boxes)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), hex_color)
        pBdr.append(el)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Inline bold / italic rendering
# ---------------------------------------------------------------------------

def _add_runs_with_inline(para, text: str, base_font_name="Calibri", base_size=Pt(11)):
    """
    Adds runs to a paragraph handling **bold** and *italic* markdown inline syntax.
    """
    # Split on bold (**...**) and italic (*...*)
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            run = para.add_run(tok[2:-2])
            run.bold = True
            run.font.name = base_font_name
            run.font.size = base_size
        elif tok.startswith("*") and tok.endswith("*"):
            run = para.add_run(tok[1:-1])
            run.italic = True
            run.font.name = base_font_name
            run.font.size = base_size
        else:
            run = para.add_run(tok)
            run.font.name = base_font_name
            run.font.size = base_size


# ---------------------------------------------------------------------------
# Block-level element writers
# ---------------------------------------------------------------------------

def add_code_block(doc: Document, lines: list[str]):
    """Render a code block: Courier New 9pt, grey background, box border."""
    code_text = "\n".join(lines)
    para = doc.add_paragraph()
    _set_para_bg(para, "F2F2F2")
    _add_full_border(para, "CCCCCC")
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.left_indent  = Cm(0.5)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_screenshot_placeholder(doc: Document, label: str):
    """Render a screenshot placeholder box: blue tint, centred label."""
    para = doc.add_paragraph()
    _set_para_bg(para, "E8F0FE")
    _add_full_border(para, "4472C4")
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[ SCREENSHOT PLACEHOLDER: {label} ]")
    run.font.name  = "Calibri"
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)


def add_blockquote(doc: Document, text: str):
    """Render a > blockquote with amber left border."""
    para = doc.add_paragraph()
    _set_para_bg(para, "FFF3CD")
    _add_para_border(para, "FFC107", border_size=6)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.8)
    _add_runs_with_inline(para, text.lstrip("> ").strip(), base_size=Pt(10.5))


def add_bullet(doc: Document, text: str, level: int = 0):
    """Add a bullet list item."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_after  = Pt(3)
    _add_runs_with_inline(para, text.lstrip("- ").strip())


def add_table(doc: Document, rows: list[list[str]]):
    """Render a markdown pipe table. First row = header."""
    if not rows:
        return
    col_count = len(rows[0])
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

            if r_idx == 0:
                # Header row — navy background, white bold text
                _set_cell_bg(cell, "1F497D")
                run = para.add_run(cell_text.strip())
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
            else:
                # Alternating row background
                bg = "FFFFFF" if r_idx % 2 == 1 else "EEF3FA"
                _set_cell_bg(cell, bg)
                _add_runs_with_inline(para, cell_text.strip(), base_size=Pt(10))

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_and_render(doc: Document, md_path: Path):
    lines  = md_path.read_text(encoding="utf-8").splitlines()
    n      = len(lines)
    i      = 0

    in_code  = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < n:
        raw = lines[i]

        # ------------------------------------------------------------------ #
        # Code block
        # ------------------------------------------------------------------ #
        if raw.strip().startswith("```"):
            if not in_code:
                flush_table()
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code  = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Screenshot placeholder
        # ------------------------------------------------------------------ #
        m = re.match(r"\[SCREENSHOT:\s*(.+?)\]", raw.strip())
        if m:
            flush_table()
            add_screenshot_placeholder(doc, m.group(1).strip())
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Horizontal rule / section divider → page break
        # ------------------------------------------------------------------ #
        if re.match(r"^-{3,}$", raw.strip()):
            flush_table()
            doc.add_page_break()
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Headings
        # ------------------------------------------------------------------ #
        if raw.startswith("#### "):
            flush_table()
            p = doc.add_paragraph(raw[5:].strip(), style="Heading 3")
            i += 1
            continue
        if raw.startswith("### "):
            flush_table()
            p = doc.add_paragraph(raw[4:].strip(), style="Heading 3")
            i += 1
            continue
        if raw.startswith("## "):
            flush_table()
            p = doc.add_paragraph(raw[3:].strip(), style="Heading 2")
            i += 1
            continue
        if raw.startswith("# "):
            flush_table()
            p = doc.add_paragraph(raw[2:].strip(), style="Heading 1")
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Table rows
        # ------------------------------------------------------------------ #
        if raw.strip().startswith("|"):
            cells = [c.strip() for c in raw.strip().strip("|").split("|")]
            # Skip separator rows (e.g., | --- | --- |)
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                i += 1
                continue
            table_buf.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # ------------------------------------------------------------------ #
        # Blockquote
        # ------------------------------------------------------------------ #
        if raw.startswith("> "):
            add_blockquote(doc, raw)
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Bullet list
        # ------------------------------------------------------------------ #
        if re.match(r"^(\s*)-\s+", raw):
            indent = len(raw) - len(raw.lstrip())
            add_bullet(doc, raw.strip(), level=indent // 2)
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Empty line — paragraph spacer
        # ------------------------------------------------------------------ #
        if raw.strip() == "":
            i += 1
            continue

        # ------------------------------------------------------------------ #
        # Normal paragraph
        # ------------------------------------------------------------------ #
        para = doc.add_paragraph()
        _add_runs_with_inline(para, raw.strip())
        i += 1

    # Flush any dangling table
    flush_table()


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover(doc: Document):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Site Intelligence Agent")
    run.font.name  = "Calibri"
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = COLOR_H1

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Technical Project Report")
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(16)
    run2.font.color.rgb = COLOR_H2

    doc.add_paragraph()

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = course.add_run("ENMGT 5400 — Applications of Artificial Intelligence\nfor Engineering Managers")
    run3.font.name = "Calibri"
    run3.font.size = Pt(13)

    doc.add_paragraph()

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = author.add_run("Nishchay Vishwanath\nCornell University, Masters of Engineering Management\nMay 2026")
    run4.font.name = "Calibri"
    run4.font.size = Pt(12)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source markdown not found: {SOURCE_MD}")

    doc = setup_document()
    add_cover(doc)
    parse_and_render(doc, SOURCE_MD)

    doc.save(OUTPUT_DOC)
    print(f"Report saved to: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
